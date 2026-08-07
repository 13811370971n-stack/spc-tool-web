"""
AI integration callbacks for all chart pages.
Provides a reusable function to add AI enhance button + download report to any page.
"""

import io
import json
from dash import callback, Input, Output, State, no_update, dcc, html
import dash_bootstrap_components as dbc

from core.llm_integration import load_config, enhance_interpretation


def ai_and_download_section(prefix: str):
    """
    Returns a layout section with AI button + Download buttons.
    Use in any page layout.
    """
    return dbc.Card([
        dbc.CardHeader("🤖 AI & 报告"),
        dbc.CardBody([
            dbc.Button("🤖 AI 增强解读", id=f"{prefix}-ai-btn", color="secondary", className="w-100 mb-2"),
            dcc.Loading(html.Div(id=f"{prefix}-ai-output", className="small mt-2"), type="dot"),
            html.Hr(),
            html.Label("下载报告:"),
            dbc.ButtonGroup([
                dbc.Button("📄 Word", id=f"{prefix}-dl-word", color="success", size="sm", outline=True),
                dbc.Button("📊 Excel", id=f"{prefix}-dl-excel", color="success", size="sm", outline=True),
            ], className="w-100"),
            dcc.Download(id=f"{prefix}-download"),
        ]),
    ], className="mt-3")


def register_ai_callback(prefix: str, chart_type: str):
    """
    Register AI enhance callback for a page.
    The page must have a Store with id=f"{prefix}-analysis-summary" containing the analysis dict.
    """

    @callback(
        Output(f"{prefix}-ai-output", "children"),
        Input(f"{prefix}-ai-btn", "n_clicks"),
        State(f"{prefix}-analysis-summary", "data"),
        prevent_initial_call=True,
    )
    def on_ai(n_clicks, summary_json):
        if not summary_json:
            return dbc.Alert("请先运行分析", color="warning", className="small py-1")

        config = load_config()
        if not config.enabled:
            return dbc.Alert("AI 未启用。请配置 ~/.spc-tool/llm_config.json", color="warning", className="small py-1")

        try:
            summary = json.loads(summary_json) if isinstance(summary_json, str) else summary_json
            result = enhance_interpretation(summary, config)
            return html.Pre(result, style={"whiteSpace": "pre-wrap", "fontSize": "12px", "maxHeight": "300px", "overflowY": "auto"})
        except Exception as e:
            return dbc.Alert(f"AI 错误: {e}", color="danger", className="small py-1")


def register_download_callback(prefix: str, chart_type: str):
    """Register Word/Excel download callbacks."""

    @callback(
        Output(f"{prefix}-download", "data"),
        Input(f"{prefix}-dl-word", "n_clicks"),
        Input(f"{prefix}-dl-excel", "n_clicks"),
        State(f"{prefix}-analysis-summary", "data"),
        prevent_initial_call=True,
    )
    def on_download(n_word, n_excel, summary_json):
        from dash import ctx
        if not summary_json:
            return no_update

        summary = json.loads(summary_json) if isinstance(summary_json, str) else summary_json
        triggered = ctx.triggered_id

        if triggered == f"{prefix}-dl-word":
            return _generate_word_download(summary, chart_type)
        elif triggered == f"{prefix}-dl-excel":
            return _generate_excel_download(summary, chart_type)
        return no_update


def _generate_word_download(summary: dict, chart_type: str):
    """Generate Word report and return as download."""
    try:
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        doc.add_heading(f"SPC Analysis Report - {chart_type}", level=0)
        doc.add_paragraph(f"Chart Type: {chart_type}")

        # Status
        in_control = summary.get("in_control", True)
        doc.add_paragraph(f"Status: {'In Control ✅' if in_control else 'Out of Control ❌'}")

        # Limits
        limits = summary.get("limits", {})
        if limits:
            doc.add_heading("Control Limits", level=1)
            for k, v in limits.items():
                doc.add_paragraph(f"{k}: {v}")

        # Violations
        violations = summary.get("violations", {})
        if violations:
            doc.add_heading("Out-of-Control Points", level=1)
            for test, points in violations.items():
                doc.add_paragraph(f"Test {test}: Subgroups {points}")

        # Stats
        doc.add_heading("Statistics", level=1)
        if "sigma_within" in summary:
            doc.add_paragraph(f"σ_within: {summary['sigma_within']:.6f}")
        if "sigma_overall" in summary:
            doc.add_paragraph(f"σ_overall: {summary['sigma_overall']:.6f}")
        if "data_summary" in summary:
            doc.add_paragraph(summary["data_summary"])

        # Capability
        cap = summary.get("capability", {})
        if cap:
            doc.add_heading("Process Capability", level=1)
            for k, v in cap.items():
                doc.add_paragraph(f"{k}: {v}")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return dcc.send_bytes(buf.getvalue(), f"SPC_{chart_type}_Report.docx")
    except ImportError:
        return no_update
    except Exception:
        return no_update


def _generate_excel_download(summary: dict, chart_type: str):
    """Generate Excel report and return as download."""
    try:
        import pandas as pd

        rows = []
        rows.append({"Item": "Chart Type", "Value": chart_type})
        rows.append({"Item": "Status", "Value": "In Control" if summary.get("in_control") else "Out of Control"})

        for k, v in summary.get("limits", {}).items():
            rows.append({"Item": k, "Value": str(v)})

        if "sigma_within" in summary:
            rows.append({"Item": "σ_within", "Value": f"{summary['sigma_within']:.6f}"})
        if "sigma_overall" in summary:
            rows.append({"Item": "σ_overall", "Value": f"{summary['sigma_overall']:.6f}"})
        if "data_summary" in summary:
            rows.append({"Item": "Data", "Value": summary["data_summary"]})

        for k, v in summary.get("capability", {}).items():
            rows.append({"Item": k, "Value": str(v)})

        violations = summary.get("violations", {})
        for test, points in violations.items():
            rows.append({"Item": f"Test {test} violations", "Value": str(points)})

        df = pd.DataFrame(rows)
        buf = io.BytesIO()
        df.to_excel(buf, index=False, sheet_name="SPC Results")
        buf.seek(0)
        return dcc.send_bytes(buf.getvalue(), f"SPC_{chart_type}_Report.xlsx")
    except Exception:
        return no_update
